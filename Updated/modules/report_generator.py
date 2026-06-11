from docx import Document
from docx.shared import Inches


class ReportGenerator:

    @staticmethod
    def create_report(
        df,
        pass_count,
        fail_count,
        not_executed_count,
        chart_path,
        output_docx
        
    ):

        total_tc = len(df)

        pass_percentage = (
            pass_count / total_tc
        ) * 100

        fail_percentage = (
            fail_count / total_tc
        ) * 100

        not_executed_testcases =(
            not_executed_count/total_tc
        ) *100

        doc = Document()

        # Title

        doc.add_heading(
            "BCM Door Functionality Validation Report",
            level=1
        )

        # Summary

        doc.add_heading(
            "Test Execution Summary",
            level=2
        )

        doc.add_paragraph(
            f"Total Test Cases : {total_tc}"
        )

        doc.add_paragraph(
            f"Passed Test Cases : {pass_count}"
        )

        doc.add_paragraph(
            f"Failed Test Cases : {fail_count}"
        )

        doc.add_paragraph(
            f"Not Executes Test Cases: {not_executed_count}"
        )

        doc.add_paragraph(
            f"Pass Percentage : {pass_percentage:.2f}%"
        )

        doc.add_paragraph(
            f"Fail Percentage : {fail_percentage:.2f}%"
        )

        doc.add_paragraph(
            f"Not_Executes_test_cases: {not_executed_testcases:.2f}%"
        
        )

        # Pie Chart

        doc.add_heading(
            "Pass / Fail Pie Chart",
            level=2
        )

        doc.add_picture(
            chart_path,
            width=Inches(4)
        )

        # Detailed Table

        doc.add_heading(
            "Detailed Test Results",
            level=2
        )

        table = doc.add_table(
            rows=1,
            cols=8
        )

        table.style = "Table Grid"

        header = table.rows[0].cells

        header[0].text = "TC_ID"
        header[1].text = "Description"
        header[2].text = "Engine Status"
        header[3].text = "Speed"
        header[4].text = "Expected"
        header[5].text = "Actual"
        header[6].text = "Status"
        header[7].text = "Remarks"

        for index, row in df.iterrows():

            cells = table.add_row().cells

            cells[0].text = str(row["TC_ID"])
            cells[1].text = str(row["Test_Description"])
            cells[2].text = str(row["Ignition_ON/OFF"])
            cells[3].text =str(row["Vehicle_Speed_(km/h)"])
            cells[4].text = str(row["Expected_Result"])
            cells[5].text = str(row["Actual_Result"])
            cells[6].text = str(row["Status"])
            cells[7].text = str(row["Remarks"])

        # Failed Cases

        doc.add_heading(
            "Failed Test Cases",
            level=2
        )

        failed_df = df[df["Status"] == "Fail"]

        if len(failed_df) == 0:

            doc.add_paragraph(
                "No Failed Test Cases."
            )

        else:

            for index, row in failed_df.iterrows():

                doc.add_paragraph(
                    f"{row['TC_ID']} : "
                    f"{row['Remarks']}"
                )

        # Conclusion

        doc.add_heading(
            "Conclusion",
            level=2
        )

        doc.add_paragraph(
            f"A total of {total_tc} test cases "
            f"were executed.\n\n"
            f"{pass_count} test cases passed.\n"
            f"{fail_count} test cases failed.\n\n"
            f"{not_executed_count} test cases not executed. \n\n"
            f"Pass Percentage : "
            f"{pass_percentage:.2f}%"
        )

        doc.save(output_docx)